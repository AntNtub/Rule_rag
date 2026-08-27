from __future__ import annotations

import json
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from pypdf import PdfReader


SUPPORTED_SUFFIXES = {".txt", ".md", ".docx", ".pdf"}


@dataclass(frozen=True)
class RegulationDocument:
    path: Path
    text: str
    title: str
    category: str
    issued_at: str | None
    source_url: str | None
    version: str


def _read_docx(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n\n".join(parts)


def _read_pdf(path: Path) -> str:
    return "\n\n".join((page.extract_text() or "") for page in PdfReader(path).pages)


def load_document(path: Path) -> RegulationDocument:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError(f"Unsupported file type: {path}")
    if suffix in {".txt", ".md"}:
        text = path.read_text(encoding="utf-8-sig")
    elif suffix == ".docx":
        text = _read_docx(path)
    else:
        text = _read_pdf(path)

    sidecar = path.with_suffix(path.suffix + ".metadata.json")
    metadata: dict[str, str] = {}
    if sidecar.exists():
        metadata = json.loads(sidecar.read_text(encoding="utf-8"))

    return RegulationDocument(
        path=path,
        text=text,
        title=metadata.get("title", path.stem),
        category=metadata.get("category", "uncategorized"),
        issued_at=metadata.get("issued_at"),
        source_url=metadata.get("source_url"),
        version=metadata.get("version", metadata.get("issued_at", "unversioned")),
    )


def discover_documents(root: Path) -> list[Path]:
    if root.is_file():
        return [root] if root.suffix.lower() in SUPPORTED_SUFFIXES else []
    return sorted(
        path
        for path in root.rglob("*")
        if path.suffix.lower() in SUPPORTED_SUFFIXES
        and path.name.lower() != "readme.md"
    )
